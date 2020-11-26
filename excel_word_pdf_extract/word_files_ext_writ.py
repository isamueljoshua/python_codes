# coding: utf-8

__author__ = ["Samuel Joshua"]

'''
I am learning to open and edit docx files

Compared to plain text files docx files are structured
These structures represent 3 different data types in python-docx

A document object represents the entire document
Document object contains Paragraph objects
Paragraph objects contains list of one or more Run objetcs
'''
import docx

d = docx.Document('demo.docx')

# all paragraphs
print("Paragraphs")
print(d.paragraphs)

print("First Paragraph")
# first paragraph
print(d.paragraphs[0])

print("Text in Para 1 and 2")
# text in first paragraph
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

p = d.paragraphs[1]
# list of run objects, A new run starts whenever there is a change in the style
# there are 4 runs here, normal Text, Bold, Normal, italic
print("Runs")
print(p.runs)
print("The Run values are: ")
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

# you can also check if it is bolded, italicised or not
print(p.runs[1].bold)
print(p.runs[0].bold == None)
print(p.runs[3].italic)

# to change the text inside doc
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'

d.save('demo2.docx')

# check the style of paragraph
print(p.style)

# changing the style of the document's text to Title style
p.style = 'Title'
d.save('demo3.docx')

# Creating a NEW DOCX document
d = docx.Document()
d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('This is another paragraph')
d.save('demo4.docx')

# Add a run to first paragraph
p = d.paragraphs[0]
p.add_run('This is a new run')
print(p.runs)
p.runs[1].bold = True
d.save('demo5.docx')


# To modify a word document (NOT implemented)
'''
Open a new word document with docx.Document and copy all paragraph and run objects to the new document
'''


# get all all text inside a word doc as a string
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))